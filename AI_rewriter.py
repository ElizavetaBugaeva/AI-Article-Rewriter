import os
import openai
from docx import Document

# Initialize the OpenAI client with an environment variable
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise ValueError("API key for OpenAI is not set. Please set the OPENAI_API_KEY environment variable.")
client = openai.OpenAI(api_key=api_key)

# Function to read an article from the Scraped_Articles directory
def read_article(file_path):
    doc = Document(file_path)
    article = "\n".join([para.text for para in doc.paragraphs])
    return article

# Function to rewrite an article using OpenAI API
def rewrite_article(article_text):
    messages = [
        {"role": "system", "content": "You are an experienced copywriter with extensive experience in writing articles about manufacturing."},
        {"role": "user", "content": (
            "Rewrite the article provided below so that it's completely original with no plagiarism. "
            "Change the order of the sections, vary the sentence length, and change the paragraph size. "
            "Do this for every section, sentence, and paragraph. "
            "Rewrite the conclusion to be about MakerVerse, an on-demand manufacturing platform.\n\n"
            f"{article_text}"
        )}
    ]

    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=messages,
        max_tokens=3000,  # Adjust max tokens based on the length of the article
        temperature=0.7
    )

    rewritten_article = response.choices[0].message.content.strip()

    return rewritten_article

# Function to save rewritten article to a new directory
def save_rewritten_article(title, content):
    sanitized_title = title.replace(" ", "_")
    file_name = f"Rewritten_Articles/{sanitized_title}.docx"
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)
    doc.save(file_name)

# Main function to process all articles
def process_all_articles():
    input_directory = 'Scrapes_Articles'
    output_directory = 'Rewritten_Articles'
    
    # Create the directories if they do not exist
    os.makedirs(input_directory, exist_ok=True)
    os.makedirs(output_directory, exist_ok=True)

    # Process all articles in the directory
    filenames = [filename for filename in os.listdir(input_directory) if filename.endswith('.docx')]
    if not filenames:
        print("No articles found in the Scraped_Articles directory.")
        return

    for filename in filenames:
        file_path = os.path.join(input_directory, filename)
        article = read_article(file_path)
        rewritten_article = rewrite_article(article)

        # Extract the title from the original document
        doc = Document(file_path)
        title = doc.paragraphs[0].text

        # Save the rewritten article
        save_rewritten_article(title, rewritten_article)
        print(f"Processed and saved: {title}")

if __name__ == "__main__":
    process_all_articles()
