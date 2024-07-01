import os
import requests
from bs4 import BeautifulSoup
import pandas as pd
from docx import Document
import re

# List of URLs to scrape
URLS = [
    "https://xometry.pro/en-eu/articles/",
    "https://xometry.pro/en-eu/articles/cnc-machining-eu/",
    "https://xometry.pro/en-eu/articles/3d-printing-eu/",
    "https://xometry.pro/en-eu/articles/sheet-metal-eu/",
    "https://xometry.pro/en-eu/articles/injection-moulding-eu/",
    "https://xometry.pro/en-eu/articles/die-casting-eu/",
    "https://xometry.pro/en-eu/articles/vacuum-casting-eu/",
    "https://xometry.pro/en-eu/articles/compression-molding-eu/",
    "https://xometry.pro/en-eu/articles/materials-eu/",
    "https://xometry.pro/en-eu/articles/design-eu/",
    "https://xometry.pro/en-eu/articles/post-processing-eu/"
]

BASE_URL = "https://xometry.pro/en-eu/articles"
OUTPUT_DIR = 'Scraped_Articles'
CSV_FILE = 'hyperlinks.csv'

# Function to get hyperlinks from a given URL
def get_hyperlinks(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    hyperlinks = {a['href'] for a in soup.find_all('a', href=True) if a['href'].startswith(BASE_URL)}
    return hyperlinks

# Function to scrape article title and content
def scrape_article(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Get the article title
    title_tag = soup.find('div', class_='article-title')
    title = title_tag.get_text(strip=True) if title_tag else 'Untitled'
    
    # Get the article content
    content_tag = soup.find('div', class_='article-content')
    content = content_tag.get_text(strip=True) if content_tag else ''
    
    return title, content

# Function to sanitize filenames
def sanitize_filename(name):
    return re.sub(r'[\\/*?:"<>|]', "_", name)

# Main function to scrape and save articles
def main():
    unique_hyperlinks = set()
    data = []

    # Scrape each URL for hyperlinks
    for url in URLS:
        hyperlinks = get_hyperlinks(url)
        for hyperlink in hyperlinks:
            if hyperlink not in unique_hyperlinks:
                data.append({'URL': url, 'Hyperlink': hyperlink})
                unique_hyperlinks.add(hyperlink)

    # Create DataFrame and save to CSV
    df = pd.DataFrame(data, columns=['URL', 'Hyperlink'])
    df.to_csv(CSV_FILE, index=False)

    # Create directory to save articles
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Scrape each hyperlink and save the article
    for hyperlink in df['Hyperlink']:
        title, content = scrape_article(hyperlink)
        if not title or not content:
            continue
        
        # Debugging statements
        print(f"Title: {title}")
        print(f"URL: {hyperlink}")
        
        # Create a new Document
        doc = Document()
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
        
        # Sanitize the filename and save the document
        sanitized_title = sanitize_filename(title)
        file_name = os.path.join(OUTPUT_DIR, f"{sanitized_title}.docx")
        doc.save(file_name)
        print(f"Saved file: {file_name}")

    print(df)

if __name__ == "__main__":
    main()
